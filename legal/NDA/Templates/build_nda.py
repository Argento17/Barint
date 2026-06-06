# -*- coding: utf-8 -*-
"""Build the Bari NDA template (one page, two columns) mirroring the source NDA structure."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Page geometry: one page, narrow margins ----
sec = doc.sections[0]
sec.top_margin = Inches(0.5)
sec.bottom_margin = Inches(0.5)
sec.left_margin = Inches(0.6)
sec.right_margin = Inches(0.6)

# Base style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(8)
pf = style.paragraph_format
pf.space_after = Pt(3)
pf.line_spacing = 1.0


def set_cols(section, num, space_twips=360):
    """Set the number of newspaper columns on a section."""
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), str(space_twips))
    cols.set(qn('w:equalWidth'), '1')


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(runs, justify=True, space_after=3, size=8):
    """runs = list of (text, bold) tuples, or a plain string."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if isinstance(runs, str):
        runs = [(runs, False)]
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
    return p


# =====================================================================
# TOP BLOCK — single column (full width) intro / parties
# =====================================================================
add_title('Nondisclosure Agreement')

add_para('This Agreement is entered into this ____ day of ___________ 20__ by '
         'Bari Ltd. and/or any Bari entity, and/or Tom Bar Haim (collectively, “Bari”), '
         'with offices at __________________________, and', size=8)

add_para([('_______________________ with offices at ____________________________, '
           'and together with Bari, the “Parties”).', False)],
         size=8)

add_para('Each party owns, possesses or has developed certain technical or business '
         'information related to its respective developments, technologies, products or '
         'product plans; and the parties hereto, desire to disclose to each other Confidential '
         'Information (as defined below) for the purpose of evaluation of a potential business '
         'relationship between the parties hereto and/or the engagement between the parties '
         'hereto in a mutually agreed upon business relationship (the ', size=8)
# Merge the Purpose sentence into one justified paragraph
last = doc.paragraphs[-1]
r = last.add_run('"Purpose"'); r.bold = True; r.font.size = Pt(8)
r = last.add_run(') and, to induce disclosure by the Parties, the Parties desire to set forth '
                 'their undertakings and obligations of confidentiality and nondisclosure.')
r.font.size = Pt(8)

add_para([('NOW, THEREFORE, the parties hereby agree as follows:', False)], size=8, space_after=4)

# =====================================================================
# COLUMN BREAK INTO TWO COLUMNS for the body
# =====================================================================
new_sec = doc.add_section(WD_SECTION.CONTINUOUS)
new_sec.top_margin = Inches(0.5)
new_sec.bottom_margin = Inches(0.5)
new_sec.left_margin = Inches(0.6)
new_sec.right_margin = Inches(0.6)
set_cols(new_sec, 2, space_twips=360)


def clause(num, heading, body):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(num + ' '); r.bold = True; r.font.size = Pt(8)
    r = p.add_run(heading + ' '); r.bold = True; r.font.size = Pt(8)
    r = p.add_run(body); r.font.size = Pt(8)
    return p


def cont(body):
    return add_para(body, size=8)


clause('1.', 'Definition.',
       '"Confidential Information" shall mean any information and data of a proprietary or '
       'confidential nature, whether in oral, written, graphic, machine-readable form, or in '
       'any form, including but not limited to technical information, technologies, developments, '
       'designs, methods, specifications, marketing, sales, price, financial information, operating, '
       'performance, cost, know-how, trade secrets, patent applications, business and process '
       'information, and all record bearing media containing or disclosing such information and '
       'techniques, whether or not marked as "Confidential", which is disclosed by one party '
       '("Discloser") to the other party ("Recipient") pursuant to this Agreement. When appropriate, '
       'the term shall also include any samples, models or prototypes, or parts thereof. The '
       'Confidential Information that disclosed by Bari hereunder shall be in the field of food and '
       'nutrition data, scoring and technology.')

cont('The confidentiality obligations of this Agreement shall not apply to any information which '
     '(a) is already in the public domain through no breach of this Agreement; (b) was, as between '
     'the parties, lawfully in Recipient’s possession prior to receipt from the Disclosing Party '
     'as is evidenced thereby, (c) is received by Recipient independently from a third party free to '
     'lawfully disclose such information to Recipient, or (d) is independently developed by Recipient '
     'without use of the Confidential Information or any portion thereof, as is evidenced thereby. '
     'Confidential Information shall not be deemed to be in the public domain merely because any part '
     'of the Confidential Information is embodied in general disclosure or because individual features, '
     'components or combinations thereof are now or become known to the public. A disclosure by '
     'Recipient of Confidential Information in response to a valid order by a court or other governmental '
     'body, or as otherwise required by law or necessary to establish the rights of either party under '
     'this Agreement, shall not be considered to be a breach of this Agreement, provided, however, that '
     'such Recipient shall provide a reasonable prior notice to Discloser, to the extent possible, to '
     'allow it to seek protective or other court orders and in any case Recipient shall limit such '
     'disclosure to the minimum disclosure required for compliance with such legal requirement.')

clause('3.', 'Restrictions.',
       'All Confidential Information delivered pursuant to this Agreement (a) shall be maintained in '
       'strict confidence, with at least the same degree of care as Recipient normally exercises to '
       'protect its own confidential information of a similar nature, but no less than a reasonable '
       'degree of care; (b) may only be disclosed to those employees and consultants of Recipient '
       '("Representatives") on a "need to know" basis for the Purpose, and provided that prior to '
       'disclosing any Confidential Information to such Representatives Recipient shall have ensured '
       'that the Representatives are aware of the provisions of this Agreement and have signed '
       'non-disclosure and non-use undertakings substantially similar to those contained in this '
       'Agreement, and provided further that Recipient shall remain liable for compliance of the '
       'Representatives with the terms hereof; (c) shall not be used by Recipient for any purpose, '
       'other than the Purpose as defined above, without the prior written consent of the Discloser; '
       'and (d) shall remain the property of and be returned to the Discloser (along with all copies '
       'thereof) within thirty (30) days of receipt by Recipient of a written request from the Discloser '
       'that sets forth the Confidential Information to be returned.')

clause('4.', 'Duration.',
       'Unless mutually agreed otherwise in writing, Recipient’s obligations hereunder with respect '
       'to each item of Confidential Information shall expire five (5) years from the date of receipt '
       'thereof by Recipient.')

clause('5.', 'Term.',
       'This Agreement shall be effective as of the date stated above and shall automatically expire '
       'five (5) years from its effective date; provided, however, that the obligations accruing prior '
       'to termination as set forth herein shall survive the termination as specified in Section 4, above.')

clause('6.', 'Mutual Disclaimers; no License.',
       'The Parties shall have no obligation to enter into any further agreement with each other. It is '
       'understood and agreed that no warranties of any kind are given by the Discloser with respect to '
       'Confidential Information provided hereunder except that the Discloser warrants it has the right to '
       'make the disclosure. Discloser shall remain the sole owner of the Confidential Information disclosed '
       'thereby hereunder. The disclosure of Confidential Information by the Discloser shall not grant the '
       'Recipient any express, implied or other license or rights to the Confidential Information or to any '
       'intellectual property of the Discloser.')

clause('7.', 'Injunctive Relief.',
       'Since a breach by Recipient of any of the promises or agreements contained herein may result in '
       'irreparable and continuing damage to Discloser for which there may be no adequate remedy at law, '
       'Discloser shall be, upon appropriate proof, entitled to seek injunctive relief and such other relief '
       'as may be proper (including monetary damages if appropriate).')

clause('8.', 'General.',
       'This Agreement represents the entire understanding and agreement of the parties and supersedes all '
       'prior communications, agreements and understandings relating to the subject matter hereof. In the '
       'event that any of the provisions of this Agreement shall be held by a court or other tribunal of '
       'competent jurisdiction to be illegal, invalid or unenforceable, such provisions shall be limited or '
       'eliminated to the minimum extent necessary so that this Agreement shall otherwise remain in full force '
       'and effect. No waiver or modification of this Agreement will be binding upon either party unless made '
       'in writing and signed by a duly authorized representative of such party and no failure or delay in '
       'enforcing any right will be deemed a waiver. This Agreement may be executed in counterparts, each of '
       'which shall be deemed an original, but both of which together shall constitute one and the same '
       'instrument. Transmission by facsimile of an executed counterpart of this Agreement shall be deemed to '
       'constitute due and sufficient delivery of such counterpart. The laws of the State of Israel shall '
       'govern this Agreement, and the competent courts of Tel Aviv, Israel shall have exclusive jurisdiction '
       'in any matter arising out of or relating to this Agreement. This Agreement may only be amended with the '
       'written consent of both Parties hereto. This Agreement may not be assigned by either party.')

# ---- Signature block ----
sigp = doc.add_paragraph()
sigp.paragraph_format.space_before = Pt(6)
r = sigp.add_run('Bari Ltd.'); r.bold = True; r.font.size = Pt(8)


def sig_line(label):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(label + '\t'); r.font.size = Pt(8)
    r = p.add_run('_______________________________'); r.font.size = Pt(8)


sig_line('By:')
sig_line('Date:')

doc.add_paragraph().paragraph_format.space_after = Pt(2)
sig_line('By:')
sig_line('Title:')
sig_line('Date:')

out = r'C:\Bari\Legal\NDA\Bari_NDA_Template.docx'
doc.save(out)
print('Saved:', out)
